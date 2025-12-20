#!/usr/bin/env python3
"""
Create Word document bug report for RT-950 Pro TNC issue.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def create_bug_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('RT-950 Pro 固件错误报告及修复建议', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run('版本: ').bold = True
    meta.add_run('V0.24\n')
    meta.add_run('日期: ').bold = True
    meta.add_run('2024年12月20日\n')
    meta.add_run('报告者: ').bold = True
    meta.add_run('业余无线电社区')
    
    # Divider
    doc.add_paragraph('─' * 50)
    
    # Bug #1
    doc.add_heading('错误 #1: TNC 类型 "KISS" 模式无法通过蓝牙发送数据 (严重)', level=1)
    
    doc.add_heading('问题描述', level=2)
    doc.add_paragraph(
        '当用户在菜单中设置 TNC Type = KISS 时，APRS数据包不会通过蓝牙发送到APRSDroid等应用程序。'
        '只有设置 TNC Type = WinAPRS 时蓝牙才能工作。\n\n'
        '这使得用户无法使用标准KISS协议连接APRSDroid。'
    )
    
    doc.add_heading('问题根因', level=2)
    doc.add_paragraph(
        '在函数 FUN_080140c0 (地址 0x080140c0) 中，代码检查 TNC 类型是否等于 1：'
    )
    
    # Code block - current buggy code
    code1 = doc.add_paragraph()
    code1.add_run('当前代码 (有问题):').bold = True
    
    code_text = '''
// 地址: 0x08014745 (反编译)
if (*(char *)(DAT_08014150 + 0x1d) == '\\x01') {  // 只检查类型 1 (WinAPRS)
    FUN_08024444(local_214, uVar3 + 1 & 0xffff);  // 发送到蓝牙
}
'''
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code_text)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    
    # TNC Type table
    doc.add_heading('TNC 类型值:', level=3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '值'
    hdr_cells[1].text = '类型名称'
    hdr_cells[2].text = '蓝牙输出'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    data = [
        ('0', 'MacAPRS', '❌ 不工作'),
        ('1', 'WinAPRS', '✅ 工作'),
        ('2', 'APRS', '❌ 不工作'),
        ('3', 'KISS', '❌ 不工作 (应该工作!)'),
    ]
    for i, (val, name, status) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = val
        row[1].text = name
        row[2].text = status
    
    doc.add_paragraph()
    
    # Fix recommendation
    doc.add_heading('修复建议', level=2)
    
    doc.add_paragraph().add_run('方案 A: 让 KISS 模式也能工作 (推荐)').bold = True
    
    fix_code = '''
// 原始代码 (FUN_080140c0, 地址约 0x08014745)
// 原始:
if (*(char *)(DAT_08014150 + 0x1d) == '\\x01') {
    FUN_08024444(local_214, uVar3 + 1 & 0xffff);
}

// 修复后:
uint8_t tnc_type = *(uint8_t *)(DAT_08014150 + 0x1d);
if (tnc_type == 1 || tnc_type == 3) {  // WinAPRS(1) 或 KISS(3)
    FUN_08024444(local_214, uVar3 + 1 & 0xffff);
}
'''
    fix_para = doc.add_paragraph()
    fix_run = fix_para.add_run(fix_code)
    fix_run.font.name = 'Consolas'
    fix_run.font.size = Pt(9)
    
    doc.add_paragraph().add_run('方案 B: 更灵活的检查 (最佳)').bold = True
    
    fix_code2 = '''
// 检查任何非零TNC类型都发送到蓝牙
uint8_t tnc_type = *(uint8_t *)(DAT_08014150 + 0x1d);
if (tnc_type > 0) {  // 任何启用的TNC模式
    FUN_08024444(local_214, uVar3 + 1 & 0xffff);
}
'''
    fix_para2 = doc.add_paragraph()
    fix_run2 = fix_para2.add_run(fix_code2)
    fix_run2.font.name = 'Consolas'
    fix_run2.font.size = Pt(9)
    
    # Assembly fix
    doc.add_heading('汇编级修复', level=2)
    
    doc.add_paragraph().add_run('原始指令 (地址 0x08014116):').bold = True
    asm1 = '''
ram:08014114    407f            ldrb        r0,[r0,#0x1d]
ram:08014116    0128            cmp         r0,#0x1         ; 只比较 1
ram:08014118    02d1            bne         LAB_08014120    ; 不等于1则跳过
'''
    asm_para1 = doc.add_paragraph()
    asm_run1 = asm_para1.add_run(asm1)
    asm_run1.font.name = 'Consolas'
    asm_run1.font.size = Pt(9)
    
    doc.add_paragraph().add_run('修复指令:').bold = True
    asm2 = '''
ram:08014114    407f            ldrb        r0,[r0,#0x1d]
ram:08014116    0028            cmp         r0,#0x0         ; 比较 0
ram:08014118    02d0            beq         LAB_08014120    ; 等于0则跳过 (任何非零都发送)
'''
    asm_para2 = doc.add_paragraph()
    asm_run2 = asm_para2.add_run(asm2)
    asm_run2.font.name = 'Consolas'
    asm_run2.font.size = Pt(9)
    
    # Binary patch section
    doc.add_heading('完整修复补丁', level=1)
    doc.add_heading('二进制补丁 (针对 v0.24)', level=2)
    
    patch_table = doc.add_table(rows=3, cols=4)
    patch_table.style = 'Table Grid'
    
    patch_hdr = patch_table.rows[0].cells
    patch_hdr[0].text = '地址'
    patch_hdr[1].text = '原始字节'
    patch_hdr[2].text = '修改字节'
    patch_hdr[3].text = '说明'
    for cell in patch_hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    patch_data = [
        ('0x08014116', '01 28', '00 28', '将 cmp r0,#0x1 改为 cmp r0,#0x0'),
        ('0x08014118', '02 D1', '02 D0', '将 bne 改为 beq'),
    ]
    for i, (addr, orig, patch, desc) in enumerate(patch_data, 1):
        row = patch_table.rows[i].cells
        row[0].text = addr
        row[1].text = orig
        row[2].text = patch
        row[3].text = desc
    
    doc.add_paragraph()
    effect = doc.add_paragraph()
    effect.add_run('效果: ').bold = True
    effect.add_run('任何非零TNC类型都会通过蓝牙发送KISS帧')
    
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run('注意: ').bold = True
    note.add_run('这些是 little-endian Thumb-2 指令')
    
    # Verification
    doc.add_heading('验证测试', level=2)
    doc.add_paragraph('修复后请测试:')
    tests = [
        '设置 TNC Type = KISS',
        '配对蓝牙并连接 APRSDroid',
        '接收APRS信号，确认APRSDroid能显示',
        '从APRSDroid发送位置，确认radio能发射',
    ]
    for test in tests:
        doc.add_paragraph(test, style='List Number')
    
    # Related functions table
    doc.add_heading('附录: 相关函数地址 (v0.24)', level=1)
    
    func_table = doc.add_table(rows=6, cols=3)
    func_table.style = 'Table Grid'
    
    func_hdr = func_table.rows[0].cells
    func_hdr[0].text = '函数'
    func_hdr[1].text = '地址'
    func_hdr[2].text = '功能'
    for cell in func_hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    func_data = [
        ('FUN_080140c0', '0x080140c0', 'KISS帧构建'),
        ('FUN_08024444', '0x08024444', 'USART1发送(蓝牙)'),
        ('FUN_0802445c', '0x0802445c', '单字节发送'),
        ('DAT_08014150', '0x2000A83C', '设置结构基址'),
        ('TNC类型偏移', '+0x1D', 'TNC类型字节'),
    ]
    for i, (func, addr, desc) in enumerate(func_data, 1):
        row = func_table.rows[i].cells
        row[0].text = func
        row[1].text = addr
        row[2].text = desc
    
    # Contact
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    contact = doc.add_heading('联系方式', level=2)
    doc.add_paragraph('如有问题请联系:')
    doc.add_paragraph('GitHub: https://github.com/Dadud/radtel-950-pro')
    doc.add_paragraph('原始项目: https://github.com/nicsure/radtel-950-pro')
    
    doc.add_paragraph()
    thanks = doc.add_paragraph()
    thanks.add_run('感谢您对产品的持续改进！').italic = True
    
    return doc


def main():
    output_dir = Path(__file__).parent.parent / 'docs'
    output_file = output_dir / 'RT950_Pro_Bug_Report_TNC_KISS.docx'
    
    print(f"Creating bug report document...")
    doc = create_bug_report()
    
    print(f"Saving to: {output_file}")
    doc.save(str(output_file))
    
    print("Done!")
    print(f"\nFile created: {output_file}")


if __name__ == '__main__':
    main()

